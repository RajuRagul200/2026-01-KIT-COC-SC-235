# =====================================
# SPSS-Equivalent Paired Samples t-Test
# T2 vs T3 Cultural Retention
# Tables + Bar Charts saved in Word
# =====================================

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from scipy.stats import ttest_rel
from docx import Document
from docx.shared import Inches

# -------------------------------
# 1. CREATE SAMPLE LONGITUDINAL DATA
# -------------------------------
np.random.seed(42)
n = 120   # sample size

data = {
    "Cultural_Awareness_T2": np.random.normal(4.35, 0.35, n),
    "Cultural_Awareness_T3": np.random.normal(4.15, 0.38, n),

    "Knowledge_Retention_T2": np.random.normal(4.40, 0.32, n),
    "Knowledge_Retention_T3": np.random.normal(4.18, 0.36, n),

    "Cultural_Attachment_T2": np.random.normal(4.30, 0.34, n),
    "Cultural_Attachment_T3": np.random.normal(4.05, 0.37, n)
}

df = pd.DataFrame(data)

# -------------------------------
# 2. CREATE WORD DOCUMENT
# -------------------------------
doc = Document()
doc.add_heading(
    "Paired Samples t-Test for Cultural Retention (T2 vs T3)",
    level=1
)

doc.add_paragraph(
    "This analysis examines whether cultural learning and attachment "
    "are sustained over time following the initial digital experience."
)

# -------------------------------
# 3. PAIRED t-TEST TABLE
# -------------------------------
doc.add_heading("Table 1. Paired Samples t-Test Results", level=2)

table = doc.add_table(rows=1, cols=6)
hdr = table.rows[0].cells
hdr[0].text = "Dependent Variable"
hdr[1].text = "Mean (T2)"
hdr[2].text = "Mean (T3)"
hdr[3].text = "t-value"
hdr[4].text = "p-value"
hdr[5].text = "Cohen’s d"

def paired_test(dv_t2, dv_t3, label):
    t, p = ttest_rel(df[dv_t2], df[dv_t3])
    diff = df[dv_t2] - df[dv_t3]
    d = diff.mean() / diff.std(ddof=1)

    row = table.add_row().cells
    row[0].text = label
    row[1].text = f"{df[dv_t2].mean():.2f}"
    row[2].text = f"{df[dv_t3].mean():.2f}"
    row[3].text = f"{t:.2f}"
    row[4].text = f"{p:.4f}"
    row[5].text = f"{d:.2f}"

paired_test(
    "Cultural_Awareness_T2",
    "Cultural_Awareness_T3",
    "Cultural Awareness"
)

paired_test(
    "Knowledge_Retention_T2",
    "Knowledge_Retention_T3",
    "Knowledge Retention"
)

paired_test(
    "Cultural_Attachment_T2",
    "Cultural_Attachment_T3",
    "Cultural Attachment"
)

# -------------------------------
# 4. BAR CHART FUNCTION (T2 vs T3)
# -------------------------------
plt.rcParams["font.family"] = "Times New Roman"
plt.rcParams["font.size"] = 14
plt.rcParams["font.weight"] = "bold"

def add_paired_bar_chart(t2, t3, ylabel, title, fig_no):
    means = [df[t2].mean(), df[t3].mean()]
    stds = [df[t2].std(), df[t3].std()]

    plt.figure(figsize=(6, 5))
    plt.bar(["T2", "T3"], means, yerr=stds, capsize=6)
    plt.ylabel(ylabel, fontweight="bold")
    plt.title(title, fontweight="bold")
    plt.ylim(1, 5)
    plt.tight_layout()

    img = f"{title.replace(' ', '_')}.png"
    plt.savefig(img, dpi=300)
    plt.close()

    doc.add_heading(f"Figure {fig_no}. {title}", level=2)
    doc.add_picture(img, width=Inches(4.8))

# -------------------------------
# 5. INSERT BAR CHARTS
# -------------------------------
add_paired_bar_chart(
    "Cultural_Awareness_T2",
    "Cultural_Awareness_T3",
    "Mean Cultural Awareness Score",
    "Cultural Awareness at T2 and T3",
    1
)

add_paired_bar_chart(
    "Knowledge_Retention_T2",
    "Knowledge_Retention_T3",
    "Mean Knowledge Retention Score",
    "Knowledge Retention at T2 and T3",
    2
)

add_paired_bar_chart(
    "Cultural_Attachment_T2",
    "Cultural_Attachment_T3",
    "Mean Cultural Attachment Score",
    "Cultural Attachment at T2 and T3",
    3
)

# -------------------------------
# 6. SAVE WORD FILE
# -------------------------------
file_name = "Paired_t_Test_Cultural_Retention_T2_T3.docx"
doc.save(file_name)

print(f"\n✅ Paired Samples t-Test results saved in Word file:\n{file_name}")

