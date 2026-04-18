# ================================
# SPSS-Equivalent One-Way ANOVA
# Tables + Multiple Bar Charts saved in Word
# ================================

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from scipy.stats import f_oneway
from statsmodels.stats.multicomp import pairwise_tukeyhsd
from docx import Document
from docx.shared import Inches

# -------------------------------
# 1. CREATE SAMPLE DATASET
# -------------------------------
np.random.seed(123)

data = {
    "Media_Type": np.repeat([1, 2, 3, 4], 30),

    "Cultural_Awareness": np.concatenate([
        np.random.normal(4.45, 0.30, 30),
        np.random.normal(4.05, 0.32, 30),
        np.random.normal(3.90, 0.35, 30),
        np.random.normal(3.70, 0.38, 30)
    ]),

    "Knowledge_Retention": np.concatenate([
        np.random.normal(4.50, 0.28, 30),
        np.random.normal(4.10, 0.30, 30),
        np.random.normal(3.95, 0.33, 30),
        np.random.normal(3.75, 0.36, 30)
    ]),

    "Cultural_Attachment": np.concatenate([
        np.random.normal(4.40, 0.29, 30),
        np.random.normal(4.00, 0.31, 30),
        np.random.normal(3.85, 0.34, 30),
        np.random.normal(3.65, 0.37, 30)
    ])
}

df = pd.DataFrame(data)

df["Media_Label"] = df["Media_Type"].map({
    1: "AR/VR",
    2: "Interactive Web",
    3: "AI Art",
    4: "3D Archive"
})

# -------------------------------
# 2. CREATE WORD DOCUMENT
# -------------------------------
doc = Document()
doc.add_heading(
    "One-Way ANOVA Results for Digital Media Type Comparison (H3)",
    level=1
)

# -------------------------------
# 3. TABLE 1: DESCRIPTIVE STATISTICS
# -------------------------------
doc.add_heading("Table 1. Descriptive Statistics (Mean ± SD)", level=2)

desc = df.groupby("Media_Label").agg(["mean", "std"]).round(2)

table1 = doc.add_table(rows=1, cols=1 + len(desc.columns))
hdr = table1.rows[0].cells
hdr[0].text = "Digital Media Type"

for i, col in enumerate(desc.columns):
    hdr[i+1].text = f"{col[0]} ({col[1]})"

for idx, row in desc.iterrows():
    row_cells = table1.add_row().cells
    row_cells[0].text = idx
    for i, val in enumerate(row):
        row_cells[i+1].text = str(val)

# -------------------------------
# 4. TABLE 2: ONE-WAY ANOVA RESULTS
# -------------------------------
doc.add_heading("Table 2. One-Way ANOVA Results", level=2)

anova_table = doc.add_table(rows=1, cols=5)
hdr = anova_table.rows[0].cells
hdr[0].text = "Dependent Variable"
hdr[1].text = "F-value"
hdr[2].text = "df"
hdr[3].text = "p-value"
hdr[4].text = "Eta Squared (η²)"

def compute_anova(dv):
    groups = [df[df["Media_Type"] == i][dv] for i in [1, 2, 3, 4]]
    F, p = f_oneway(*groups)

    grand_mean = df[dv].mean()
    ss_between = sum(len(g) * (g.mean() - grand_mean) ** 2 for g in groups)
    ss_total = sum((df[dv] - grand_mean) ** 2)
    eta_sq = ss_between / ss_total

    return round(F, 2), "3, 116", f"{p:.4f}", round(eta_sq, 3)

for dv in ["Cultural_Awareness", "Knowledge_Retention", "Cultural_Attachment"]:
    F, df_val, p, eta = compute_anova(dv)
    row = anova_table.add_row().cells
    row[0].text = dv.replace("_", " ")
    row[1].text = str(F)
    row[2].text = df_val
    row[3].text = p
    row[4].text = str(eta)

# -------------------------------
# 5. BAR CHART FUNCTION
# -------------------------------
plt.rcParams["font.family"] = "Times New Roman"
plt.rcParams["font.size"] = 14
plt.rcParams["font.weight"] = "bold"

def add_bar_chart(df, dv, ylabel, figure_title, fig_number):
    means = df.groupby("Media_Label")[dv].mean()
    stds = df.groupby("Media_Label")[dv].std()

    plt.figure(figsize=(8, 6))
    plt.bar(means.index, means.values, yerr=stds.values, capsize=6)
    plt.xlabel("Digital Media Type", fontweight="bold")
    plt.ylabel(ylabel, fontweight="bold")
    plt.title(figure_title, fontweight="bold")
    plt.ylim(1, 5)
    plt.tight_layout()

    img_name = f"{dv}_Bar_Chart.png"
    plt.savefig(img_name, dpi=300)
    plt.close()

    doc.add_heading(f"Figure {fig_number}. {figure_title}", level=2)
    doc.add_picture(img_name, width=Inches(5.5))

# -------------------------------
# 6. INSERT BAR CHARTS
# -------------------------------
add_bar_chart(
    df,
    "Cultural_Awareness",
    "Mean Cultural Awareness Score",
    "Cultural Awareness Across Digital Media Types",
    1
)

add_bar_chart(
    df,
    "Knowledge_Retention",
    "Mean Knowledge Retention Score",
    "Knowledge Retention Across Digital Media Types",
    2
)

add_bar_chart(
    df,
    "Cultural_Attachment",
    "Mean Cultural Attachment Score",
    "Cultural Attachment Across Digital Media Types",
    3
)

# -------------------------------
# 7. SAVE WORD FILE
# -------------------------------
file_name = "One_Way_ANOVA_Digital_Media_Results_With_All_Charts.docx"
doc.save(file_name)

print(f"\n✅ Results + ALL Bar Charts saved in Word file:\n{file_name}")
