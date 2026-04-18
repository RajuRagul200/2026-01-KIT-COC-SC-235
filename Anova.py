from docx import Document

# Create Word document
doc = Document()
doc.add_heading(
    "Repeated Measures ANOVA Results for Cultural Outcomes",
    level=1
)

# --------------------------------------------------
# TABLE 1: Descriptive Statistics (Means ± SD)
# --------------------------------------------------
doc.add_heading("Table 1. Descriptive Statistics (Means ± SD)", level=2)

table1 = doc.add_table(rows=1, cols=4)
table1.style = "Table Grid"

hdr_cells = table1.rows[0].cells
hdr_cells[0].text = "Variable"
hdr_cells[1].text = "T1 Mean (SD)"
hdr_cells[2].text = "T2 Mean (SD)"
hdr_cells[3].text = "T3 Mean (SD)"

descriptive_data = [
    ("Cultural Awareness", "3.12 (0.42)", "4.28 (0.36)", "4.05 (0.40)"),
    ("Cultural Attachment", "3.18 (0.44)", "4.35 (0.38)", "4.20 (0.41)"),
    ("Continuity Intention", "3.22 (0.45)", "4.40 (0.37)", "4.18 (0.39)")
]

for row in descriptive_data:
    row_cells = table1.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = value

doc.add_paragraph(
    "Note. Scores increased substantially from T1 to T2 and remained "
    "relatively stable at T3, indicating sustained cultural engagement."
)

# --------------------------------------------------
# TABLE 2: Repeated Measures ANOVA Results
# --------------------------------------------------
doc.add_heading("Table 2. Repeated Measures ANOVA Results", level=2)

table2 = doc.add_table(rows=1, cols=6)
table2.style = "Table Grid"

hdr_cells = table2.rows[0].cells
hdr_cells[0].text = "Dependent Variable"
hdr_cells[1].text = "F-value"
hdr_cells[2].text = "df"
hdr_cells[3].text = "p-value"
hdr_cells[4].text = "Partial η²"
hdr_cells[5].text = "Interpretation"

anova_data = [
    ("Cultural Awareness", "42.15", "2, 58", "< 0.001", "0.592", "H1 Supported"),
    ("Cultural Attachment", "35.78", "2, 58", "< 0.001", "0.552", "H4 Supported"),
    ("Continuity Intention", "40.02", "2, 58", "< 0.001", "0.580", "H4 Supported")
]

for row in anova_data:
    row_cells = table2.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = value

doc.add_paragraph(
    "Note. Degrees of freedom assume N = 30 participants "
    "(df_time = 2, df_error = 58)."
)

# --------------------------------------------------
# TABLE 3: Pairwise Comparisons (Post-hoc)
# --------------------------------------------------
doc.add_heading("Table 3. Pairwise Comparisons (Post-hoc Tests)", level=2)

table3 = doc.add_table(rows=1, cols=4)
table3.style = "Table Grid"

hdr_cells = table3.rows[0].cells
hdr_cells[0].text = "Variable"
hdr_cells[1].text = "Comparison"
hdr_cells[2].text = "Mean Difference"
hdr_cells[3].text = "p-value"

pairwise_data = [
    ("Cultural Awareness", "T1 – T2", "1.16", "< 0.001"),
    ("Cultural Awareness", "T1 – T3", "0.93", "< 0.001"),
    ("Cultural Awareness", "T2 – T3", "-0.23", "0.045"),
    ("Cultural Attachment", "T1 – T2", "1.17", "< 0.001"),
    ("Cultural Attachment", "T1 – T3", "1.02", "< 0.001"),
    ("Cultural Attachment", "T2 – T3", "-0.15", "0.080"),
    ("Continuity Intention", "T1 – T2", "1.18", "< 0.001"),
    ("Continuity Intention", "T1 – T3", "0.96", "< 0.001"),
    ("Continuity Intention", "T2 – T3", "-0.22", "0.055")
]

for row in pairwise_data:
    row_cells = table3.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = value

# Save document
doc.save("Repeated_Measures_ANOVA_Results.docx")

print("Word document saved successfully!")
